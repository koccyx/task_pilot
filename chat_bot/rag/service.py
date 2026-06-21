"""RAG document ingestion and search service."""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Optional
from uuid import NAMESPACE_URL, uuid4, uuid5

from .embedding import TextEmbedder, build_embedder
from .repository import (
    PostgresRagRepository,
    RagRepository,
    StoredChunk,
    StoredDocument,
)
from .retrieval import bm25_search, reciprocal_rank_fusion
from .settings import RagSettings
from .vector_store import QdrantVectorStore, VectorSearchResult, VectorStore


@dataclass(frozen=True)
class IngestedDocument:
    """Result of a document upload."""

    document: StoredDocument
    chunk_count: int


class RagService:
    """Coordinate document storage, chunking and vector indexing."""

    def __init__(
        self,
        settings: RagSettings,
        repository: Optional[RagRepository] = None,
        vector_store: Optional[VectorStore] = None,
        embedder: Optional[TextEmbedder] = None,
    ) -> None:
        self.settings = settings
        self.repository = repository or PostgresRagRepository.from_env()
        self.embedder = embedder or build_embedder(
            provider=settings.embedding_provider,
            dimension=settings.embedding_dimension,
            model=settings.embedding_model,
            base_url=settings.embedding_base_url,
            api_key=settings.embedding_api_key,
            timeout_seconds=settings.embedding_timeout_seconds,
        )
        self.vector_store = vector_store or QdrantVectorStore(
            url=settings.qdrant_url,
            collection_name=settings.qdrant_collection,
            vector_size=settings.embedding_dimension,
        )

    async def ingest_file(
        self,
        filename: str,
        content_type: str,
        content: bytes,
    ) -> IngestedDocument:
        """Store an uploaded file and index its extracted text."""
        if not filename:
            raise ValueError("Имя файла не указано")
        if not content:
            raise ValueError("Файл пустой")

        document_id = uuid4().hex
        text = self.extract_text(filename=filename, content=content)
        chunks_text = self.chunk_text(text)
        if not chunks_text:
            raise ValueError("Не удалось извлечь текст из файла")

        document = StoredDocument(
            id=document_id,
            filename=filename,
            content_type=content_type or "application/octet-stream",
            size_bytes=len(content),
            stored_path="postgres://rag_documents.content",
            uploaded_at=PostgresRagRepository.now_iso(),
        )
        chunks = [
            StoredChunk(
                id=str(uuid5(NAMESPACE_URL, f"{document_id}-{index}")),
                document_id=document_id,
                chunk_index=index,
                text=chunk_text,
            )
            for index, chunk_text in enumerate(chunks_text)
        ]

        await self.repository.add_document(document, chunks, content)
        vectors = self._embed_documents([chunk.text for chunk in chunks])
        self.vector_store.upsert_chunks(
            [
                {
                    "chunk_id": chunk.id,
                    "document_id": document.id,
                    "filename": document.filename,
                    "text": chunk.text,
                    "vector": vectors[index],
                }
                for index, chunk in enumerate(chunks)
            ]
        )
        return IngestedDocument(document=document, chunk_count=len(chunks))

    async def list_documents(self) -> list[StoredDocument]:
        """Return uploaded documents."""
        return await self.repository.list_documents()

    async def get_document_with_chunks(
        self,
        document_id: str,
    ) -> tuple[StoredDocument, list[StoredChunk]]:
        """Return document and extracted chunks."""
        document = await self.repository.get_document(document_id)
        if document is None:
            raise KeyError(document_id)
        return document, await self.repository.list_chunks(document_id)

    async def delete_document(self, document_id: str) -> bool:
        """Delete a document, all its chunks and indexed vectors."""
        document = await self.repository.get_document(document_id)
        if document is None:
            return False
        chunks = await self.repository.list_chunks(document_id)
        self.vector_store.delete_chunks([chunk.id for chunk in chunks])
        return await self.repository.delete_document(document_id)

    async def delete_chunk(self, chunk_id: str) -> StoredChunk | None:
        """Delete one chunk and its indexed vector."""
        chunk = await self.repository.get_chunk(chunk_id)
        if chunk is None:
            return None
        self.vector_store.delete_chunks([chunk.id])
        deleted = await self.repository.delete_chunk(chunk.id)
        return chunk if deleted else None

    async def search(self, query: str, limit: int = 5) -> list[VectorSearchResult]:
        """Search indexed chunks with dense retrieval, BM25 and RRF."""
        if not query.strip():
            return []
        candidate_limit = max(limit * 4, limit)
        dense_results = self.vector_store.search(
            self.embedder.embed(query),
            limit=candidate_limit,
        )
        chunks = await self.repository.list_all_chunks()
        lexical_results = bm25_search(query, chunks, limit=candidate_limit)
        fused_results = reciprocal_rank_fusion(
            dense_results=dense_results,
            lexical_results=lexical_results,
            limit=limit,
        )
        if any(not result.filename for result in fused_results):
            documents = await self.repository.list_documents()
            filename_by_document_id = {
                document.id: document.filename for document in documents
            }
            fused_results = [
                VectorSearchResult(
                    chunk_id=result.chunk_id,
                    document_id=result.document_id,
                    filename=result.filename
                    or filename_by_document_id.get(result.document_id, ""),
                    text=result.text,
                    score=result.score,
                )
                for result in fused_results
            ]
        return fused_results

    def _embed_documents(self, texts: list[str]) -> list[list[float]]:
        embed_documents = getattr(self.embedder, "embed_documents", None)
        if callable(embed_documents):
            return embed_documents(texts)
        return [self.embedder.embed_document(text) for text in texts]

    async def chunk_count(self, document_id: str) -> int:
        """Return number of chunks for a document."""
        return await self.repository.count_chunks(document_id)

    def extract_text(self, filename: str, content: bytes) -> str:
        """Extract plain text from supported document types."""
        suffix = Path(filename).suffix.lower()
        if suffix == ".pdf":
            return self._extract_pdf(content)
        if suffix == ".docx":
            return self._extract_docx(content)
        return content.decode("utf-8", errors="ignore")

    def chunk_text(self, text: str) -> list[str]:
        """Split text into overlapping token chunks."""
        normalized = re.sub(r"\s+", " ", text).strip()
        if not normalized:
            return []

        chunk_size = self.settings.chunk_size
        overlap = min(self.settings.chunk_overlap, chunk_size - 1)
        token_matches = list(re.finditer(r"\S+", normalized))
        if not token_matches:
            return []

        chunks: list[str] = []
        start = 0
        while start < len(token_matches):
            end = min(start + chunk_size, len(token_matches))
            chunk_start = token_matches[start].start()
            chunk_end = token_matches[end - 1].end()
            chunks.append(normalized[chunk_start:chunk_end].strip())
            if end == len(token_matches):
                break
            start = end - overlap
        return chunks

    @staticmethod
    def _extract_pdf(content: bytes) -> str:
        from io import BytesIO

        from pypdf import PdfReader

        reader = PdfReader(BytesIO(content))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    @staticmethod
    def _extract_docx(content: bytes) -> str:
        from io import BytesIO

        from docx import Document

        document = Document(BytesIO(content))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)
